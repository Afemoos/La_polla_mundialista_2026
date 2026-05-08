"""Genera los archivos .docx a partir de los .md"""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path, title):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_path, 'r') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if line.startswith('# ') and not line.startswith('## '):
            h = doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=2)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 40)
        elif line.startswith('| ') and ' | ' in line:
            # Table
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                row = [c.strip() for c in lines[i].split('|')[1:-1]]
                rows.append(row)
                i += 1
            i -= 1
            if len(rows) > 1 and all(c.strip() == '---' or '--' in c for c in rows[1]):
                rows.pop(1)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        table.rows[ri].cells[ci].text = cell
        elif line.startswith('- '):
            # List
            while i < len(lines) and lines[i].startswith('- '):
                text = lines[i][2:]
                # Bold markers
                text = re.sub(r'\*\*(.*?)\*\*', lambda m: m.group(1), text)
                p = doc.add_paragraph(text, style='List Bullet')
                i += 1
            i -= 1
        elif line.startswith('> '):
            text = line[2:]
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(1)
            p.runs[0].italic = True if p.runs else False
        elif line.strip() == '':
            pass
        elif line.strip():
            # Regular paragraph with bold support
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', lambda m: m.group(1), text)
            p = doc.add_paragraph(text)
        
        i += 1
    
    doc.save(docx_path)
    print(f"✅ {docx_path}")

# Generar ambos
md_to_docx(
    'Informe-Detallado_Funcionamiento_05-07-26.md',
    'Informe-Detallado_Funcionamiento_05-07-26.docx',
    'Informe Detallado de Funcionamiento'
)
md_to_docx(
    'La-Polla-Mundialista_2026.md',
    'La-Polla-Mundialista_2026.docx',
    'La Polla Mundialista 2026'
)
